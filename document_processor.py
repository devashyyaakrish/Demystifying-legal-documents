import os
import io
import PyPDF2
import docx
import torch
import time
from typing import Tuple, Dict, Any, Optional
from google.cloud import vision
from docx import Document
from transformers import AutoModelForCausalLM, AutoTokenizer
from config import Config

class DocumentProcessor:
    def __init__(self):
        self.config = Config()
        self.model = None
        self.tokenizer = None
        self.initialized = False
        self.error_message = None
        
        # Initialize the text simplification model
        print("Initializing text simplification model...")
        try:
            if self.config.USE_LOCAL_MODEL:
                print(f"Loading local model: {self.config.LOCAL_MODEL_NAME}")
                self.tokenizer = AutoTokenizer.from_pretrained(
                    self.config.LOCAL_MODEL_NAME,
                    trust_remote_code=True
                )
                self.model = AutoModelForCausalLM.from_pretrained(
                    self.config.LOCAL_MODEL_NAME,
                    torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                    trust_remote_code=True
                )
                if torch.cuda.is_available():
                    self.model = self.model.to('cuda')
                print(f"Successfully loaded local model: {self.config.LOCAL_MODEL_NAME}")
                self.initialized = True
            else:
                # Use Gemini API
                import google.generativeai as genai
                from google.api_core.exceptions import InvalidArgument
                
                if not self.config.GEMINI_API_KEY:
                    self.error_message = "Error: GEMINI_API_KEY is not set in the .env file"
                    print(self.error_message)
                    return
                    
                try:
                    genai.configure(api_key=self.config.GEMINI_API_KEY)
                    # Test the API key with a simple request
                    self.model = genai.GenerativeModel(
                        model_name=self.config.MODEL_NAME,
                        generation_config=self.config.GENERATION_CONFIG,
                        safety_settings=self.config.SAFETY_SETTINGS
                    )
                    print(f"Successfully initialized Gemini API with model: {self.config.MODEL_NAME}")
                    self.initialized = True
                except Exception as e:
                    self.error_message = f"Error initializing Gemini API: {str(e)}"
                    print(self.error_message)
                
        except Exception as e:
            self.error_message = f"Error initializing model: {str(e)}"
            print(self.error_message)
        
        # Initialize Google Cloud Vision client
        self.vision_client = vision.ImageAnnotatorClient()
        
    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """
        Analyze a document and extract/process its text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dict containing the extracted/processed text and metadata
        """
        try:
            # Get file extension
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            
            # Extract text based on file type
            if ext == '.pdf':
                text = self._extract_text_from_pdf(file_path)
            elif ext in ['.docx', '.doc']:
                text = self._extract_text_from_docx(file_path)
            elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.webp']:
                text = self._extract_text_with_ocr(file_path)
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                return {
                    'success': False,
                    'error': f'Unsupported file type: {ext}'
                }
            
            # Check if we got any text
            if not text or not text.strip():
                return {
                    'success': False,
                    'error': 'No text could be extracted from the document',
                    'file_type': ext
                }
            
            # Simplify the extracted text
            simplified_result = self.simplify_text(text)
            
            if not simplified_result['success']:
                return simplified_result
                
            return {
                'success': True,
                'original_text': text,
                'simplified_text': simplified_result['simplified_text'],
                'char_count': len(text),
                'word_count': len(text.split()),
                'file_type': ext
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': f'Error processing document: {str(e)}'
            }
            
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from a PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            # If text extraction fails, try OCR
            print(f"PDF text extraction failed, trying OCR: {str(e)}")
            return self._extract_text_with_ocr(file_path)
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = docx.Document(file_path)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
    
    def _extract_text_with_ocr(self, image_path: str) -> str:
        """Extract text from an image using Google Cloud Vision OCR."""
        try:
            with open(image_path, 'rb') as image_file:
                content = image_file.read()
                
            image = vision.Image(content=content)
            response = self.vision_client.text_detection(image=image)
            
            if response.error.message:
                raise Exception(f'Error from Vision API: {response.error.message}')
                
            if response.text_annotations:
                return response.text_annotations[0].description
                
            return ""
            
        except Exception as e:
            raise Exception(f"OCR processing failed: {str(e)}")

    def extract_text(self, file_path: str) -> Tuple[bool, str]:
        """Extract text from various document formats including scanned PDFs and images."""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            if file_ext == '.pdf':
                # First try to extract text directly
                success, text = self._extract_text_from_pdf(file_path)
                if success and text.strip():
                    return success, text
                
                # If direct extraction fails or returns empty, try OCR
                print("Direct text extraction failed, trying OCR...")
                return self._ocr_pdf(file_path)
                
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp', '.webp']:
                return self._ocr_image(file_path)
                
            elif file_ext == '.docx':
                return self._extract_from_docx(file_path)
                
            elif file_ext == '.txt':
                return self._extract_from_txt(file_path)
                
            else:
                return False, f"Unsupported file format: {file_ext}"
                
        except Exception as e:
            return False, f"Error extracting text: {str(e)}"
    
    def _extract_text_from_pdf(self, file_path: str) -> Tuple[bool, str]:
        """Extract text from searchable PDFs."""
        try:
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                text = '\n'.join([page.extract_text() for page in reader.pages])
                return True, text.strip()
        except Exception as e:
            return False, f"Error reading PDF: {str(e)}"
    
    def _ocr_pdf(self, file_path: str) -> Tuple[bool, str]:
        """Perform OCR on PDF using Google Cloud Vision."""
        try:
            with open(file_path, 'rb') as file:
                content = file.read()
            
            image = vision.Image(content=content)
            response = self.vision_client.document_text_detection(image=image)
            
            if not response.text_annotations:
                return False, "No text found in the PDF"
                
            return True, response.text_annotations[0].description
            
        except Exception as e:
            return False, f"Error during PDF OCR: {str(e)}"
    
    def _ocr_image(self, file_path: str) -> Tuple[bool, str]:
        """Perform OCR on image files using Google Cloud Vision."""
        try:
            with open(file_path, 'rb') as file:
                content = file.read()
            
            image = vision.Image(content=content)
            response = self.vision_client.document_text_detection(image=image)
            
            if not response.text_annotations:
                return False, "No text found in the image"
                
            return True, response.text_annotations[0].description
            
        except Exception as e:
            return False, f"Error during image OCR: {str(e)}"

    def _extract_from_docx(self, file_path: str) -> Tuple[bool, str]:
        """Extract text from DOCX files."""
        try:
            doc = Document(file_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            return True, text.strip()
        except Exception as e:
            return False, f"Error reading DOCX: {str(e)}"

    def _extract_from_txt(self, file_path: str) -> Tuple[bool, str]:
        """Extract text from TXT files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                return True, text.strip()
        except Exception as e:
            return False, f"Error reading text file: {str(e)}"

    def _process_text_in_chunks(self, text: str, chunk_size: int = None) -> Dict[str, Any]:
        """Process long text in chunks to handle large inputs efficiently.
        
        Args:
            text: The input text to process
            chunk_size: Optional custom chunk size in characters. If None, uses config.CHUNK_SIZE
            
        Returns:
            Dict containing the processed text and metadata
        """
        config = Config()
        chunk_size = chunk_size or config.CHUNK_SIZE
        max_retries = 3
        start_time = time.time()
        
        print(f"Processing text of length: {len(text)} characters")
        
        # Split text into paragraphs for better chunking
        paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
        chunks = []
        current_chunk = []
        current_length = 0
        
        # Build chunks of appropriate size
        for para in paragraphs:
            para_length = len(para)
            
            # If adding this paragraph would exceed chunk size, finalize current chunk
            if current_length + para_length > chunk_size and current_chunk:
                chunks.append('\n\n'.join(current_chunk))
                current_chunk = []
                current_length = 0
                
            current_chunk.append(para)
            current_length += para_length
        
        # Add the last chunk if not empty
        if current_chunk:
            chunks.append('\n\n'.join(current_chunk))
        
        print(f"Split text into {len(chunks)} chunks for processing")
        
        # Process chunks with retries
        simplified_chunks = []
        
        for i, chunk in enumerate(chunks, 1):
            print(f"Processing chunk {i}/{len(chunks)} ({(i/len(chunks))*100:.1f}%)")
            
            # Add a small delay between API calls to avoid rate limiting
            if i > 1:
                time.sleep(0.3)
            
            # Retry logic for each chunk
            for attempt in range(max_retries):
                try:
                    result = self.simplify_text(chunk)
                    
                    if result['success']:
                        simplified_chunks.append(result['simplified_text'])
                        break
                    elif attempt < max_retries - 1:
                        print(f"Attempt {attempt + 1} failed, retrying...")
                        time.sleep(1)  # Wait before retry
                    else:
                        return {
                            'success': False,
                            'error': f"Failed to process chunk {i} after {max_retries} attempts: {result.get('error', 'Unknown error')}",
                            'original_length': len(text)
                        }
                        
                except Exception as e:
                    if attempt < max_retries - 1:
                        print(f"Error processing chunk {i}, attempt {attempt + 1}: {str(e)}")
                        time.sleep(1)  # Wait before retry
                    else:
                        return {
                            'success': False,
                            'error': f"Error processing chunk {i}: {str(e)}",
                            'original_length': len(text)
                        }
        
        # Combine results with proper spacing
        combined_text = '\n\n'.join(simplified_chunks)
        
        total_time = time.time() - start_time
        print(f"Processed {len(chunks)} chunks in {total_time:.2f} seconds "
              f"({len(text)/total_time:.1f} chars/sec)")
              
        return {
            'success': True,
            'simplified_text': combined_text,
            'original_length': len(text),
            'simplified_length': len(combined_text),
        }
        
    # Check if model is properly initialized
    if not self.initialized:
        error_msg = self.error_message or "Model not properly initialized"
        print(f"Error: {error_msg}")
        return {
            'success': False,
            'error': error_msg,
            'original_length': len(text)
        }
        
    try:
        # Check if we should process in chunks
        if len(text) > self.config.CHUNK_SIZE:
            print(f"Text length ({len(text)} chars) exceeds chunk size, processing in chunks...")
            return self._process_text_in_chunks(text)
        
        # Optimized prompt for better quality and clarity
        prompt = """
        Please simplify the following legal text to make it more understandable while preserving all key information.
        
        Guidelines:
        1. Use simpler words and shorter sentences
        2. Break down complex legal terms and explain them in parentheses
        3. Maintain the original meaning and legal accuracy
        4. Keep important legal terms but explain them
        5. Use bullet points or numbered lists where appropriate
        6. Keep the structure clear with proper paragraphs
        
        Text to simplify:
        """ + text
        
        # Send request with optimized parameters
        response = self.model.generate_content(
            prompt,
            generation_config={
                **self.config.GENERATION_CONFIG,
                'max_output_tokens': min(
                    self.config.GENERATION_CONFIG['max_output_tokens'],
                    len(text) * 2  # Allow longer output if needed
                )
            }
        )
        
        # Handle response
        if not response or not hasattr(response, 'text') or not response.text:
            error_msg = "No valid response from the model"
            print(f"Error: {error_msg}")
            return {
                'success': False,
                'error': error_msg,
                'original_length': len(text)
            }
        
        # Clean up the response
        simplified_text = response.text.strip()
        
        # If the response seems truncated, try with a smaller chunk
        if len(simplified_text) < len(text) / 10:  # If response is less than 10% of original
            print("Response seems too short, trying with smaller chunks...")
            return self._process_text_in_chunks(text, chunk_size=self.config.CHUNK_SIZE // 2)
        
        return {
            'success': True,
            'simplified_text': simplified_text,
            'original_length': len(text),
            'simplified_length': len(simplified_text),
            'model': f"gemini:{self.config.MODEL_NAME}",
            'processed_in_chunks': False
        }
        
    except Exception as e:
        error_msg = str(e)
        print(f"Error in simplify_text: {error_msg}")
        
        # If the error is related to length, try with smaller chunks
        if any(err in error_msg.lower() for err in ["too many tokens", "maximum context length", "index out of range"]):
            print("Text too long, trying with smaller chunks...")
            return self._process_text_in_chunks(text, chunk_size=self.config.CHUNK_SIZE // 2)
            
        return {
            'success': False,
            'error': f'Error: {error_msg[:500]}',  # Truncate long error messages
            'original_length': len(text)
        }
                    'max_output_tokens': min(
                        self.config.GENERATION_CONFIG['max_output_tokens'],
                        len(text) * 2  # Allow longer output if needed
                    )
                }
            )
            
            # Handle response
            if not response or not hasattr(response, 'text') or not response.text:
                error_msg = "No valid response from the model"
                print(f"Error: {error_msg}")
                return {
                    'success': False,
                    'error': error_msg,
                    'original_length': len(text)
                }
            
            # Clean up the response
            simplified_text = response.text.strip()
            
            # If the response seems truncated, try with a smaller chunk
            if len(simplified_text) < len(text) / 10:  # If response is less than 10% of original
                print("Response seems too short, trying with smaller chunks...")
                return self._process_text_in_chunks(text, chunk_size=self.config.CHUNK_SIZE // 2)
            
            return {
                'success': True,
                'simplified_text': simplified_text,
                'original_length': len(text),
                'simplified_length': len(simplified_text),
                'model': f"gemini:{self.config.MODEL_NAME}",
                'processed_in_chunks': False
            }
            
        except Exception as e:
            error_msg = str(e)
            print(f"Error in simplify_text: {error_msg}")
            
            # If the error is related to length, try with smaller chunks
            if any(err in error_msg.lower() for err in ["too many tokens", "maximum context length", "index out of range"]):
                print("Text too long, trying with smaller chunks...")
                return self._process_text_in_chunks(text, chunk_size=self.config.CHUNK_SIZE // 2)
                
            return {
                'success': False,
                'error': f'Error: {error_msg[:500]}',  # Truncate long error messages
                'original_length': len(text)
            }
